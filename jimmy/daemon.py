"""
Jimmy background daemon — scheduler + inbox watcher + daily pre-compiler.

Runs as a thread inside the FastAPI server process.
- Scheduler: runs ingestion jobs on configurable intervals
- Inbox watcher: monitors ~/jimmy/inbox/ for new files
- Daily compiler: pre-generates digest/spark/daily at 5am
"""
import json
import logging
import threading
import time
from datetime import datetime, timedelta
from pathlib import Path

logger = logging.getLogger("jimmy.daemon")

# ── Default sync schedule (hours between runs) ─────────────────────────────────
DEFAULT_SCHEDULE = {
    "apple_notes": 6,
    "google_calendar": 2,
    "gmail": 6,
    "spotify": 24,
    "canvas": 12,
    "notion": 12,
    "readwise": 24,
    "pocket": 24,
    "trakt": 24,
    "gdrive": 12,
    "granola": 6,
    "github": 24,
    "whoop": 24,
}

SCHEDULE_PATH = Path.home() / ".jimmy" / "sync_schedule.json"
LOG_PATH = Path.home() / ".jimmy" / "sync.log"
INBOX_DIR = Path.home() / "jimmy" / "inbox"
PROCESSED_DIR = INBOX_DIR / "processed"


def _setup_logging():
    LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    handler = logging.FileHandler(str(LOG_PATH), encoding="utf-8")
    handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
    logger.addHandler(handler)
    logger.setLevel(logging.INFO)


def _load_schedule() -> dict:
    if SCHEDULE_PATH.exists():
        try:
            custom = json.loads(SCHEDULE_PATH.read_text())
            merged = {**DEFAULT_SCHEDULE}
            merged.update(custom)
            return merged
        except Exception:
            pass
    return dict(DEFAULT_SCHEDULE)


def _save_default_schedule():
    SCHEDULE_PATH.parent.mkdir(parents=True, exist_ok=True)
    if not SCHEDULE_PATH.exists():
        SCHEDULE_PATH.write_text(json.dumps(DEFAULT_SCHEDULE, indent=2))


# ── Ingestion runner ───────────────────────────────────────────────────────────

def _run_source(source: str) -> dict:
    """Run a single source ingestion. Returns {ok, chunks, error}."""
    from .config import (
        CANVAS_API_TOKEN, CANVAS_API_URL,
        NOTION_API_TOKEN, READWISE_API_TOKEN,
        POCKET_CONSUMER_KEY, POCKET_ACCESS_TOKEN,
        TRAKT_CLIENT_ID, TRAKT_USERNAME,
        SPOTIFY_CLIENT_ID, SPOTIFY_CLIENT_SECRET,
        GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET,
        WHOOP_CLIENT_ID, WHOOP_CLIENT_SECRET,
        GITHUB_TOKEN,
        CHROMA_DIR,
    )
    from .storage.store import JimmyStore
    from .cli import chunk_text

    store = JimmyStore(CHROMA_DIR)

    def _chunk_and_store(docs):
        from .ingestion.base import Document
        chunks, metadatas, ids = [], [], []
        seen = set()
        ingested_at = datetime.utcnow().isoformat()
        for doc in docs:
            prefix = f"[{doc.source.upper()}: {doc.title}]\n\n"
            for i, chunk in enumerate(chunk_text(doc.content)):
                cid = f"{doc.id}_c{i}"
                if cid not in seen:
                    seen.add(cid)
                    chunks.append(prefix + chunk)
                    meta = {**doc.metadata}
                    meta.setdefault("created_at", ingested_at)
                    meta.setdefault("ingested_at", ingested_at)
                    meta["title"] = doc.title
                    meta["source"] = doc.source
                    metadatas.append(meta)
                    ids.append(cid)
        if chunks:
            store.upsert(chunks, metadatas, ids)
        return len(chunks)

    try:
        if source == "apple_notes":
            from .ingestion.apple_notes import AppleNotesIngester
            docs = AppleNotesIngester().ingest()
        elif source == "google_calendar" and GOOGLE_CLIENT_ID:
            from .ingestion.google_auth import get_all_credentials
            from .ingestion.google_calendar import GoogleCalendarIngester
            accounts = get_all_credentials(GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET)
            docs = []
            for label, creds in accounts:
                docs.extend(GoogleCalendarIngester(creds, label).ingest())
        elif source == "gmail" and GOOGLE_CLIENT_ID:
            from .ingestion.google_auth import get_all_credentials
            from .ingestion.gmail import GmailIngester
            accounts = get_all_credentials(GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET)
            docs = []
            for label, creds in accounts:
                docs.extend(GmailIngester(creds, label).ingest(days=30))
        elif source == "spotify" and SPOTIFY_CLIENT_ID:
            from .ingestion.spotify import SpotifyIngester
            docs = SpotifyIngester(SPOTIFY_CLIENT_ID, SPOTIFY_CLIENT_SECRET).ingest()
        elif source == "canvas" and CANVAS_API_TOKEN:
            from .ingestion.canvas import CanvasIngester
            docs = CanvasIngester(CANVAS_API_TOKEN, CANVAS_API_URL).ingest()
        elif source == "notion" and NOTION_API_TOKEN:
            from .ingestion.notion import NotionIngester
            docs = NotionIngester(NOTION_API_TOKEN).ingest()
        elif source == "readwise" and READWISE_API_TOKEN:
            from .ingestion.readwise import ReadwiseIngester
            docs = ReadwiseIngester(READWISE_API_TOKEN).ingest()
        elif source == "pocket" and POCKET_CONSUMER_KEY:
            from .ingestion.pocket import PocketIngester
            docs = PocketIngester(POCKET_CONSUMER_KEY, POCKET_ACCESS_TOKEN).ingest()
        elif source == "trakt" and TRAKT_CLIENT_ID:
            from .ingestion.trakt import TraktIngester
            docs = TraktIngester(TRAKT_CLIENT_ID, TRAKT_USERNAME).ingest()
        elif source == "gdrive" and GOOGLE_CLIENT_ID:
            from .ingestion.google_auth import get_all_credentials
            from .ingestion.google_drive import GoogleDriveIngester
            accounts = get_all_credentials(GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET)
            docs = []
            for label, creds in accounts:
                docs.extend(GoogleDriveIngester(creds, label).ingest())
        elif source == "granola":
            from .ingestion.granola import GranolaIngester
            docs = GranolaIngester().ingest()
        elif source == "github" and GITHUB_TOKEN:
            from .ingestion.github_repos import GitHubReposIngester
            docs = GitHubReposIngester(GITHUB_TOKEN).ingest()
        elif source == "whoop" and WHOOP_CLIENT_ID:
            from .ingestion.whoop import WhoopIngester
            docs = WhoopIngester(WHOOP_CLIENT_ID, WHOOP_CLIENT_SECRET).ingest(days=30)
        else:
            return {"ok": False, "error": f"No credentials or unknown source: {source}"}

        n = _chunk_and_store(docs)
        return {"ok": True, "chunks": n, "documents": len(docs)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ── Inbox watcher ──────────────────────────────────────────────────────────────

def _ingest_file(path: Path):
    """Auto-detect file type and ingest."""
    from .ingestion.base import Document, _h
    from .cli import chunk_text
    from .storage.store import JimmyStore
    from .config import CHROMA_DIR

    suffix = path.suffix.lower()
    content = ""
    title = path.stem

    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            content = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif suffix in {".txt", ".md", ".csv", ".log"}:
            content = path.read_text(encoding="utf-8", errors="replace")
        elif suffix == ".json":
            # Try to detect known formats (ChatGPT, Claude exports)
            raw = json.loads(path.read_text(encoding="utf-8", errors="replace"))
            if isinstance(raw, list) and raw and "mapping" in raw[0]:
                # ChatGPT export
                from .ingestion.chatgpt_export import ChatGPTExportIngester
                docs = ChatGPTExportIngester().ingest(str(path))
                _store_docs_direct(docs)
                return True
            elif isinstance(raw, list) and raw and "chat_messages" in raw[0]:
                # Claude export
                from .ingestion.claude_export import ClaudeExportIngester
                docs = ClaudeExportIngester().ingest(str(path))
                _store_docs_direct(docs)
                return True
            else:
                content = json.dumps(raw, indent=2)[:10000]
        elif suffix in {".png", ".jpg", ".jpeg", ".heic", ".webp", ".tiff"}:
            # OCR via Claude Vision
            from .ingestion.document_ocr import DocumentOCRIngester
            docs = DocumentOCRIngester().ingest(str(path))
            _store_docs_direct(docs)
            return True
        elif suffix == ".xml":
            # Check if Apple Health export
            head = path.read_text(encoding="utf-8", errors="replace")[:500]
            if "HealthData" in head or "HKQuantityTypeIdentifier" in head:
                from .ingestion.apple_health import AppleHealthIngester
                docs = AppleHealthIngester().ingest(str(path))
                _store_docs_direct(docs)
                return True
            content = head[:5000]
        else:
            content = path.read_text(encoding="utf-8", errors="replace")[:5000]

        if not content or len(content.strip()) < 30:
            return False

        doc = Document(
            id=f"inbox_{_h(str(path))}",
            content=content,
            source="file",
            title=title,
            metadata={"type": "inbox", "original_path": str(path)},
        )
        _store_docs_direct([doc])
        return True

    except Exception as e:
        logger.error(f"Inbox ingest failed for {path.name}: {e}")
        return False


def _store_docs_direct(docs):
    """Store documents directly (bypass CLI console output)."""
    from .cli import chunk_text, is_low_quality_chunk
    from .storage.store import JimmyStore
    from .config import CHROMA_DIR

    store = JimmyStore(CHROMA_DIR)
    chunks, metadatas, ids = [], [], []
    seen = set()
    ingested_at = datetime.utcnow().isoformat()
    for doc in docs:
        prefix = f"[{doc.source.upper()}: {doc.title}]\n\n"
        for i, chunk in enumerate(chunk_text(doc.content)):
            cid = f"{doc.id}_c{i}"
            if cid not in seen:
                if is_low_quality_chunk(chunk):
                    continue
                seen.add(cid)
                chunks.append(prefix + chunk)
                meta = {**doc.metadata}
                meta.setdefault("created_at", ingested_at)
                meta.setdefault("ingested_at", ingested_at)
                meta["title"] = doc.title
                meta["source"] = doc.source
                metadatas.append(meta)
                ids.append(cid)
    if chunks:
        store.upsert(chunks, metadatas, ids)
    return len(chunks)


class InboxWatcher:
    """Watch ~/jimmy/inbox/ for new files, ingest them, move to processed/."""

    def __init__(self):
        INBOX_DIR.mkdir(parents=True, exist_ok=True)
        PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
        self._observer = None

    def start(self):
        try:
            from watchdog.observers import Observer
            from watchdog.events import FileSystemEventHandler
        except ImportError:
            logger.warning("watchdog not installed — inbox watcher disabled. pip install watchdog")
            return

        watcher = self

        class Handler(FileSystemEventHandler):
            def on_created(self, event):
                if event.is_directory:
                    return
                path = Path(event.src_path)
                if path.parent == PROCESSED_DIR:
                    return
                # Small delay to let file finish writing
                time.sleep(1)
                try:
                    logger.info(f"Inbox: new file detected — {path.name}")
                    ok = _ingest_file(path)
                    if ok:
                        dest = PROCESSED_DIR / path.name
                        # Handle name conflicts
                        if dest.exists():
                            stem = path.stem
                            suffix = path.suffix
                            dest = PROCESSED_DIR / f"{stem}_{int(time.time())}{suffix}"
                        path.rename(dest)
                        logger.info(f"Inbox: ingested and moved — {path.name}")
                    else:
                        logger.warning(f"Inbox: could not ingest — {path.name}")
                except Exception as e:
                    logger.error(f"Inbox: error processing {path.name}: {e}")

        self._observer = Observer()
        self._observer.schedule(Handler(), str(INBOX_DIR), recursive=False)
        self._observer.daemon = True
        self._observer.start()
        logger.info(f"Inbox watcher started — monitoring {INBOX_DIR}")

    def stop(self):
        if self._observer:
            self._observer.stop()
            self._observer.join(timeout=5)


# ── Daily pre-compiler ─────────────────────────────────────────────────────────

def compile_daily():
    """Pre-generate digest, daily extras, spark, and news so /today is instant."""
    from .config import JIMMY_DATA_DIR
    logger.info("Daily compile: starting pre-generation...")

    try:
        from .retrieval.engine import JimmyEngine
        engine = JimmyEngine()

        # Daily extras (fact, vocab, motivational_note)
        try:
            result = engine.daily_extras()
            cache_path = JIMMY_DATA_DIR / "daily_cache.json"
            result["cached_at"] = datetime.now().isoformat()
            cache_path.write_text(json.dumps(result))
            logger.info("Daily compile: daily extras cached")
        except Exception as e:
            logger.error(f"Daily compile: daily extras failed — {e}")

        # Digest
        try:
            result = engine.digest()
            cache_path = JIMMY_DATA_DIR / "digest_cache.json"
            result["cached_at"] = datetime.now().isoformat()
            cache_path.write_text(json.dumps(result))
            logger.info("Daily compile: digest cached")
        except Exception as e:
            logger.error(f"Daily compile: digest failed — {e}")

        # Spark
        try:
            result = engine.spark()
            cache_path = JIMMY_DATA_DIR / "sparks_cache.json"
            result["cached_at"] = datetime.now().isoformat()
            cache_path.write_text(json.dumps(result))
            logger.info("Daily compile: sparks cached")
        except Exception as e:
            logger.error(f"Daily compile: sparks failed — {e}")

        logger.info("Daily compile: done")
    except Exception as e:
        logger.error(f"Daily compile: fatal error — {e}")


# ── Scheduler ──────────────────────────────────────────────────────────────────

class SyncScheduler:
    """Background thread that runs ingestion jobs on schedule."""

    def __init__(self):
        self._thread = None
        self._stop_event = threading.Event()
        self._last_run: dict[str, float] = {}
        self._last_compile: float = 0

    def start(self):
        _save_default_schedule()
        self._thread = threading.Thread(target=self._loop, daemon=True, name="jimmy-scheduler")
        self._thread.start()
        logger.info("Sync scheduler started")

    def stop(self):
        self._stop_event.set()
        if self._thread:
            self._thread.join(timeout=10)

    def _loop(self):
        # Wait 60s after startup before first sync (let server warm up)
        self._stop_event.wait(60)

        while not self._stop_event.is_set():
            try:
                self._tick()
            except Exception as e:
                logger.error(f"Scheduler tick error: {e}")
            # Check every 5 minutes
            self._stop_event.wait(300)

    def _tick(self):
        schedule = _load_schedule()
        now = time.time()

        for source, interval_hours in schedule.items():
            last = self._last_run.get(source, 0)
            if now - last >= interval_hours * 3600:
                logger.info(f"Scheduler: running {source}")
                result = _run_source(source)
                self._last_run[source] = now
                if result.get("ok"):
                    logger.info(f"Scheduler: {source} done — {result.get('chunks', 0)} chunks")
                else:
                    logger.warning(f"Scheduler: {source} failed — {result.get('error', 'unknown')}")

        # Daily compile at 5am (or if never run)
        now_dt = datetime.now()
        if now_dt.hour == 5 and (now - self._last_compile > 3600):
            self._last_compile = now
            try:
                compile_daily()
            except Exception as e:
                logger.error(f"Daily compile failed: {e}")


# ── Main daemon entry point ────────────────────────────────────────────────────

_scheduler: SyncScheduler | None = None
_inbox_watcher: InboxWatcher | None = None


def start_daemon():
    """Start all background services. Called from server startup."""
    global _scheduler, _inbox_watcher
    _setup_logging()
    logger.info("Jimmy daemon starting...")

    # Scheduler
    _scheduler = SyncScheduler()
    _scheduler.start()

    # Inbox watcher
    _inbox_watcher = InboxWatcher()
    _inbox_watcher.start()

    logger.info("Jimmy daemon running")


def stop_daemon():
    """Stop all background services."""
    global _scheduler, _inbox_watcher
    if _scheduler:
        _scheduler.stop()
    if _inbox_watcher:
        _inbox_watcher.stop()
    logger.info("Jimmy daemon stopped")
